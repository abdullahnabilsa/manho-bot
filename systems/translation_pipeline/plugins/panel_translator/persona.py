# systems/translation_pipeline/plugins/panel_translator/persona.py
from __future__ import annotations

import io
import logging
import os
import sys
from typing import Any, Dict, List, ClassVar, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from pydantic import ValidationError

from systems.translation_pipeline.base_persona import BasePersona
from systems.translation_pipeline.models.panel_data import PanelElement, Panel, PanelPageData, TranslationMetadata
from systems.translation_pipeline.models.page_job import PageJob
from systems.translation_pipeline.models.page_data import PageData
from utils.markdown_escaper import escape_markdown_v2
from utils.regex_template_engine import RegexTemplateEngine

logger = logging.getLogger(__name__)

ALLOWED_PANEL_ELEMENT_KEYS = {"type", "character", "original_text", "translated_text", "description"}

TYPE_DISPLAY_MAP = {
    "speech_bubble": "# (فقاعة عادية)",
    "background_speech": "$ (كلام بالخلفية)",
    "sfx": "& (مؤثرات صوتية)",
    "side_bubble": "* (كلام ب طرف الفقاعة)"
}

def get_display_type(elem_type: Optional[str]) -> str:
    if not elem_type: return ""
    return TYPE_DISPLAY_MAP.get(elem_type, elem_type)

class SafePanelContext:
    def __init__(self, panel: Panel) -> None:
        self.panel_index = panel.panel_index

class SafePanelElementContext:
    def __init__(self, elem: PanelElement) -> None:
        self.display_type = escape_markdown_v2(get_display_type(elem.type)) if elem.type else None
        self.character = escape_markdown_v2(elem.character) if elem.character else None
        self.original_text = escape_markdown_v2(elem.original_text) if elem.original_text else None
        self.translated_text = escape_markdown_v2(elem.translated_text) if elem.translated_text else None
        self.description = escape_markdown_v2(elem.description) if elem.description else None

class PanelMessageBuilder:
    MAX_LENGTH = 3500
    HEADER_TEMPLATE = (
        "📖 ترجمة المانهوا\n"
        "📄 الملف: {file_name}\n"
        "الرسالة: {msg_num} من [[TOTAL_MSGS]]\n"
        "━━━━━━━━━━━━━━━\n"
    )
    CONTINUATION_MARKER = "Panel \\(Continued\\)\n\n"
    FOOTER_NEXT = "\n\n━━━━━━━━━━━━━━━\nانتهى الجزء\nيتبع\\.\\.\\."
    FOOTER_END = "\n\n━━━━━━━━━━━━━━━\nاكتملت ترجمة الصفحة\\."

    _panel_template: ClassVar[str] = ""
    _element_template: ClassVar[str] = ""
    _templates_loaded: ClassVar[bool] = False

    @classmethod
    def _load_templates(cls) -> None:
        if not cls._templates_loaded:
            module = sys.modules[PanelPersona.__module__]
            base_dir = os.path.dirname(os.path.abspath(module.__file__))
            panel_path = os.path.join(base_dir, "templates", "panel_header_template.txt")
            elem_path = os.path.join(base_dir, "templates", "panel_element_layout_template.txt")
            
            with open(panel_path, "r", encoding="utf-8") as f:
                cls._panel_template = f.read()
            with open(elem_path, "r", encoding="utf-8") as f:
                cls._element_template = f.read()
            cls._templates_loaded = True

    def __init__(self, msg_num: int, is_continuation: bool, file_name: str = "Unknown") -> None:
        self._load_templates()
        self.msg_num = msg_num
        self.is_continuation = is_continuation
        self._buffer = ""
        self.file_name = escape_markdown_v2(file_name) if file_name else "Unknown"

    def _get_header(self) -> str:
        return self.HEADER_TEMPLATE.format(file_name=self.file_name, msg_num=self.msg_num)

    def _get_footer(self, is_last_message: bool) -> str:
        return self.FOOTER_END if is_last_message else self.FOOTER_NEXT

    def format_panel_header(self, panel: Panel) -> str:
        ctx = SafePanelContext(panel)
        engine = RegexTemplateEngine(self._panel_template)
        return engine.render(ctx) + "\n"

    def format_element(self, elem: PanelElement) -> str:
        ctx = SafePanelElementContext(elem)
        engine = RegexTemplateEngine(self._element_template)
        return engine.render(ctx) + "\n"

    def can_fit(self, text_to_add: str) -> bool:
        if not text_to_add: return True
        projected_length = len(self._get_header()) + len(self.CONTINUATION_MARKER if self.is_continuation else "") + len(self._buffer) + len(text_to_add) + len(self.FOOTER_NEXT)
        return projected_length <= self.MAX_LENGTH

    def append_to_buffer(self, text: str) -> None:
        if text:
            self._buffer += text

    def build(self, is_last_message: bool) -> str:
        header = self._get_header()
        continuation = self.CONTINUATION_MARKER if self.is_continuation else ""
        footer = self._get_footer(is_last_message)
        return f"{header}{continuation}{self._buffer}{footer}"

class PanelPersona(BasePersona):
    name = "Panel Translator"

    async def validate_and_update_job(self, job: PageJob, raw_json: Dict[str, Any]) -> PageJob:
        job_id = job.job_id
        logger.info(f"JobID={job_id} | Starting Panel JSON validation")

        try:
            raw_meta = raw_json.get("translation_metadata", {})
            if not isinstance(raw_meta, dict): raw_meta = {}
            metadata = TranslationMetadata(
                source_language=raw_meta.get("source_language"),
                target_language=raw_meta.get("target_language"),
                style=raw_meta.get("style")
            )

            raw_panels = raw_json.get("panels", [])
            if not isinstance(raw_panels, list): raw_panels = []

            sanitized_panels: List[Panel] = []

            for p_idx, raw_panel in enumerate(raw_panels):
                try:
                    if not isinstance(raw_panel, dict): continue
                    panel_index = raw_panel.get("panel_index", p_idx + 1)
                    
                    raw_elements = raw_panel.get("elements", [])
                    if not isinstance(raw_elements, list): raw_elements = []

                    sanitized_elements: List[PanelElement] = []
                    for e_idx, raw_elem in enumerate(raw_elements):
                        try:
                            clean_dict = self._sanitize_element(raw_elem, job_id)
                            element = PanelElement(**clean_dict)
                            sanitized_elements.append(element)
                        except ValidationError as ve:
                            salvaged = self._salvage_element(clean_dict, ve)
                            if salvaged: sanitized_elements.append(salvaged)

                    panel = Panel(panel_index=panel_index, elements=sanitized_elements)
                    sanitized_panels.append(panel)

                except Exception as panel_e:
                    logger.warning(f"JobID={job_id} | Panel {p_idx} failed: {panel_e}")

            panel_data = PanelPageData(translation_metadata=metadata, panels=sanitized_panels, file_name=job.file_name)
            
            if not job.page_data:
                job.page_data = PageData()
            job.page_data.custom_data = panel_data
            
            logger.info(f"JobID={job_id} | Panel Validation successful. Panels: {len(panel_data.panels)}")
            return job

        except Exception as e:
            logger.error(f"JobID={job_id} | Critical Panel validation failure: {e}", exc_info=True)
            if not job.page_data:
                job.page_data = PageData()
            return job

    def _sanitize_element(self, raw_elem: Any, job_id: Any) -> Dict[str, Any]:
        if not isinstance(raw_elem, dict): return {}
        clean_dict: Dict[str, Any] = {}
        for key, value in raw_elem.items():
            lower_key = key.lower()
            if lower_key in ALLOWED_PANEL_ELEMENT_KEYS:
                clean_dict[lower_key] = value if value != "" else None
        return clean_dict

    def _salvage_element(self, clean_dict: Dict[str, Any], ve: ValidationError) -> PanelElement:
        salvage_dict = clean_dict.copy()
        for error in ve.errors():
            loc = error.get("loc", [])
            if loc and isinstance(loc, tuple) and len(loc) > 0:
                field = loc[0]
                if field in ALLOWED_PANEL_ELEMENT_KEYS:
                    salvage_dict[field] = None
        try: return PanelElement(**salvage_dict)
        except Exception: return PanelElement()

    async def paginate(self, job: PageJob, mode: str = "scene_split") -> List[str]:
        file_name = job.file_name or "Unknown"
        
        if not job.page_data or not job.page_data.custom_data or not isinstance(job.page_data.custom_data, PanelPageData):
            builder = PanelMessageBuilder(msg_num=1, is_continuation=False, file_name=file_name)
            return [builder.build(is_last_message=True).replace("[[TOTAL_MSGS]]", "1")]

        panel_data: PanelPageData = job.page_data.custom_data
        panels = panel_data.panels
        
        messages: List[str] = []
        current_builder = PanelMessageBuilder(msg_num=1, is_continuation=False, file_name=file_name)

        for p_idx, panel in enumerate(panels):
            is_last_panel = (p_idx == len(panels) - 1)
            panel_header = current_builder.format_panel_header(panel)
            
            if not current_builder.can_fit(panel_header):
                messages.append(current_builder.build(is_last_message=False))
                current_builder = PanelMessageBuilder(msg_num=len(messages)+1, is_continuation=(mode == "single_message"), file_name=file_name)
            
            current_builder.append_to_buffer(panel_header)

            for elem in panel.elements:
                formatted_elem = current_builder.format_element(elem)
                
                if current_builder.can_fit(formatted_elem):
                    current_builder.append_to_buffer(formatted_elem)
                else:
                    messages.append(current_builder.build(is_last_message=False))
                    current_builder = PanelMessageBuilder(msg_num=len(messages)+1, is_continuation=True, file_name=file_name)
                    
                    if panel.panel_index:
                        cont_header = f"━━━━ 🖼️ *Panel {panel.panel_index}* \\(متابعة\\) ━━━━\n\n"
                        current_builder.append_to_buffer(cont_header)
                    
                    current_builder.append_to_buffer(formatted_elem)

            if mode == "scene_split" and not is_last_panel:
                if current_builder._buffer.strip():
                    messages.append(current_builder.build(is_last_message=False))
                    current_builder = PanelMessageBuilder(msg_num=len(messages)+1, is_continuation=False, file_name=file_name)

        messages.append(current_builder.build(is_last_message=True))
        total_msgs = str(len(messages))
        return [msg.replace("[[TOTAL_MSGS]]", total_msgs) for msg in messages]

    def generate_txt(self, pages: List[PageData]) -> io.BytesIO:
        buffer = io.StringIO()
        
        for page_idx, page_data in enumerate(pages, 1):
            file_name = page_data.file_name or "Unknown"
            buffer.write("═" * 60 + "\n")
            buffer.write(f"  📄 Page {page_idx} | 🖼️ File: {file_name}\n")
            buffer.write("═" * 60 + "\n\n")
            
            panel_data = page_data.custom_data if isinstance(page_data.custom_data, PanelPageData) else None
            if not panel_data:
                buffer.write("  ⚠️ No panel data extracted.\n\n")
                continue
                
            for panel in panel_data.panels:
                buffer.write(f"🖼️ Panel {panel.panel_index}\n")
                buffer.write("-" * 60 + "\n")
                
                for elem in panel.elements:
                    elem_type = elem.type or "text"
                    display_type = get_display_type(elem_type)
                    
                    buffer.write(f"  ┌── Element ({display_type})\n")
                    if elem.character:
                        buffer.write(f"  │ 🗣 Character: {elem.character}\n")
                    if elem.original_text:
                        buffer.write(f"  │ 📝 Original: {elem.original_text}\n")
                    if elem.translated_text:
                        buffer.write(f"  │ 🇸🇦 Translation: {elem.translated_text}\n")
                    if elem.description:
                        buffer.write(f"  │ 💡 Description: {elem.description}\n")
                    buffer.write("  └──────────────────────────────────────────\n")
                buffer.write("\n")
                
        val = buffer.getvalue()
        buffer.close()
        return io.BytesIO(val.encode('utf-8'))

    def generate_docx(self, pages: List[PageData]) -> io.BytesIO:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        for page_idx, page_data in enumerate(pages, 1):
            file_name = page_data.file_name or "Unknown"
            
            p_title = doc.add_paragraph()
            run_title = p_title.add_run(f'📄 Page {page_idx}')
            run_title.bold = True
            run_title.font.size = Pt(18)
            run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            
            p_file = doc.add_paragraph()
            run_file = p_file.add_run(f'🖼️ File: {file_name}')
            run_file.italic = True
            run_file.font.size = Pt(10)
            run_file.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            
            p_div = doc.add_paragraph()
            p_div.add_run('_' * 60).font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)
            
            panel_data = page_data.custom_data if isinstance(page_data.custom_data, PanelPageData) else None
            if not panel_data:
                doc.add_paragraph("⚠️ No panel data extracted.")
                doc.add_page_break()
                continue

            for panel in panel_data.panels:
                h = doc.add_heading(f'🖼️ Panel {panel.panel_index}', level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x2E, 0x74, 0x05)
                    
                for elem in panel.elements:
                    elem_type = elem.type or "text"
                    display_type = get_display_type(elem_type)
                    
                    p_elem = doc.add_paragraph()
                    p_elem.paragraph_format.left_indent = Pt(20)
                    
                    run_type = p_elem.add_run(f'({display_type})')
                    run_type.italic = True
                    run_type.font.size = Pt(9)
                    run_type.font.color.rgb = RGBColor(0xA6, 0xA6, 0xA6)
                    
                    if elem.character:
                        run_spk = p_elem.add_run(f' - 🗣 {elem.character}')
                        run_spk.bold = True
                        run_spk.font.size = Pt(11)
                        run_spk.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        
                    if elem.original_text:
                        p_orig = doc.add_paragraph()
                        p_orig.paragraph_format.left_indent = Pt(40)
                        run_orig_lbl = p_orig.add_run('📝 Original: ')
                        run_orig_lbl.bold = True
                        p_orig.add_run(elem.original_text)
                        
                    if elem.translated_text:
                        p_trans = doc.add_paragraph()
                        p_trans.paragraph_format.left_indent = Pt(40)
                        run_trans_lbl = p_trans.add_run('🇸🇦 Translation: ')
                        run_trans_lbl.bold = True
                        run_trans_lbl.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                        run_trans_text = p_trans.add_run(elem.translated_text)
                        run_trans_text.font.size = Pt(12)
                        
                    if elem.description:
                        p_desc = doc.add_paragraph()
                        p_desc.paragraph_format.left_indent = Pt(40)
                        run_desc_lbl = p_desc.add_run('💡 Description: ')
                        run_desc_lbl.bold = True
                        p_desc.add_run(elem.description).italic = True
                        
                    doc.add_paragraph()
                doc.add_paragraph()
                
            doc.add_page_break()

        byte_io = io.BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        return byte_io