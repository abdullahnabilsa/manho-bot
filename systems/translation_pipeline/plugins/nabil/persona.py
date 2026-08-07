# systems/translation_pipeline/plugins/nabil/persona.py
from __future__ import annotations

import io
import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor

from systems.translation_pipeline.base_persona import BasePersona
from systems.translation_pipeline.models.script_data import ScriptPageData
from systems.translation_pipeline.models.page_job import PageJob
from systems.translation_pipeline.models.page_data import PageData
from utils.markdown_escaper import escape_markdown_v2

logger = logging.getLogger(__name__)

class NabilPersona(BasePersona):
    name = "NABIL"

    async def validate_and_update_job(self, job: PageJob, raw_json: Dict[str, Any]) -> PageJob:
        job_id = job.job_id
        logger.info(f"JobID={job_id} | Starting Script JSON validation")

        try:
            script_text = raw_json.get("script", "")
            if not isinstance(script_text, str):
                script_text = str(script_text)
            
            script_text = script_text.strip()
            script_data = ScriptPageData(script=script_text, file_name=job.file_name)
            
            if not job.page_data:
                job.page_data = PageData()
            job.page_data.custom_data = script_data
            
            logger.info(f"JobID={job_id} | Script Validation successful. Length: {len(script_text)}")
            return job

        except Exception as e:
            logger.error(f"JobID={job_id} | Critical Script validation failure: {e}", exc_info=True)
            if not job.page_data:
                job.page_data = PageData()
            return job

    async def paginate(self, job: PageJob, mode: str = "scene_split") -> List[str]:
        file_name = job.file_name or "Unknown"
        escaped_file_name = escape_markdown_v2(file_name)
        
        if not job.page_data or not job.page_data.custom_data or not isinstance(job.page_data.custom_data, ScriptPageData):
            text = f"📖 ترجمة المانهوا\n📄 الملف: {escaped_file_name}\n━━━━━━━━━━━━━━━\n⚠️ لا توجد بيانات."
            return [text]
            
        script_data: ScriptPageData = job.page_data.custom_data
        raw_lines = script_data.script.split('\n')
        
        messages: List[str] = []
        header = (
            f"📖 ترجمة المانهوا\n"
            f"📄 الملف: {escaped_file_name}\n"
            f"━━━━━━━━━━━━━━━\n"
        )
        footer = "\n\n━━━━━━━━━━━━━━━\nاكتملت ترجمة الصفحة\\."
        
        current_msg_lines = []
        current_length = len(header) + len(footer)
        
        for line in raw_lines:
            escaped_line = escape_markdown_v2(line)
            line_len = len(escaped_line) + 1
            
            if current_length + line_len > 3500:
                msg_text = header + "\n".join(current_msg_lines) + footer
                messages.append(msg_text)
                current_msg_lines = [escaped_line]
                current_length = len(header) + len(footer) + line_len
            else:
                current_msg_lines.append(escaped_line)
                current_length += line_len
                
        if current_msg_lines:
            msg_text = header + "\n".join(current_msg_lines) + footer
            messages.append(msg_text)
            
        total_msgs = str(len(messages))
        return [msg.replace("[[TOTAL_MSGS]]", total_msgs) for msg in messages]

    def generate_txt(self, pages: List[PageData], session_note: Optional[str] = None) -> io.BytesIO:
        buffer = io.StringIO()
        
        if session_note:
            buffer.write("═" * 60 + "\n")
            buffer.write("  📝 ملاحظة الجلسة:\n")
            buffer.write(f"  {session_note}\n")
            buffer.write("═" * 60 + "\n\n")
        
        buffer.write("═" * 60 + "\n")
        buffer.write("  💡 ملاحظة الرموز:\n")
        buffer.write("  # = فقاعة عادية\n")
        buffer.write("  $ = كلام بالخلفية\n")
        buffer.write("  & = مؤثرات\n")
        buffer.write("  * = كلام ب طرف الفقاعة\n")
        buffer.write("═" * 60 + "\n\n")
        
        for page_idx, page_data in enumerate(pages, 1):
            file_name = page_data.file_name or "Unknown"
            buffer.write("═" * 60 + "\n")
            buffer.write(f"  📄 Page {page_idx} | 🖼️ File: {file_name}\n")
            buffer.write("═" * 60 + "\n\n")
            
            script_data = page_data.custom_data if isinstance(page_data.custom_data, ScriptPageData) else None
            if not script_data or not script_data.script:
                buffer.write("  ⚠️ No script data extracted.\n\n")
                continue
                
            buffer.write(script_data.script + "\n\n")
                
        val = buffer.getvalue()
        buffer.close()
        return io.BytesIO(val.encode('utf-8'))

    def generate_docx(self, pages: List[PageData], session_note: Optional[str] = None) -> io.BytesIO:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

        if session_note:
            p_note = doc.add_paragraph()
            run_note = p_note.add_run(f"📝 ملاحظة الجلسة: {session_note}")
            run_note.italic = True
            run_note.font.size = Pt(12)
            run_note.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            doc.add_paragraph()

        p_legend = doc.add_paragraph()
        run_legend = p_legend.add_run("💡 ملاحظة الرموز:\n# = فقاعة عادية | $ = كلام بالخلفية | & = مؤثرات | * = كلام ب طرف الفقاعة")
        run_legend.italic = True
        run_legend.font.size = Pt(10)
        run_legend.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        doc.add_paragraph()

        for page_idx, page_data in enumerate(pages, 1):
            file_name = page_data.file_name or "Unknown"
            
            p_title = doc.add_paragraph()
            run_title = p_title.add_run(f'📄 Page {page_idx}')
            run_title.bold = True
            run_title.font.size = Pt(18)
            run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            
            p_file = doc.add_paragraph()
            run_file = p_file.add_run(f'🖼️ File: {file_name}')
            run_file.italic = True
            run_file.font.size = Pt(10)
            run_file.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            
            p_div = doc.add_paragraph()
            p_div.add_run('_' * 60).font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)
            
            script_data = page_data.custom_data if isinstance(page_data.custom_data, ScriptPageData) else None
            if not script_data or not script_data.script:
                doc.add_paragraph("⚠️ No script data extracted.")
                doc.add_page_break()
                continue

            lines = script_data.script.split('\n')
            for line in lines:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                
            doc.add_page_break()

        byte_io = io.BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        return byte_io