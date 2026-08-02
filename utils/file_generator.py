# utils/file_generator.py
from __future__ import annotations

import io
from typing import List
from docx import Document
from docx.shared import Pt, RGBColor

# Forward reference for type hinting
from systems.translation_pipeline.models.page_data import PageData

class FileGenerator:
    @staticmethod
    def generate_txt(page_data_list: List[PageData]) -> io.BytesIO:
        buffer = io.StringIO()
        for page_idx, page_data in enumerate(page_data_list, 1):
            file_name = page_data.file_name or "Unknown"
            buffer.write("═" * 60 + "\n")
            buffer.write(f"  📄 Page {page_idx} | 🖼️ File: {file_name}\n")
            buffer.write("═" * 60 + "\n\n")
            if not page_data or not page_data.scenes:
                buffer.write("  ⚠️ No data extracted.\n\n")
                continue
            for scene in page_data.scenes:
                env = scene.environment or "Unknown Environment"
                buffer.write(f"🎬 Scene {scene.scene_number}: {env}\n")
                buffer.write("-" * 60 + "\n")
                for elem in scene.elements:
                    speaker = elem.speaker or "N/A"
                    elem_type = elem.type or "text"
                    buffer.write(f"  ┌── Element {elem.element_number} ({elem_type})\n")
                    buffer.write(f"  │ 🗣 Speaker: {speaker}\n")
                    if elem.original: buffer.write(f"  │ 📝 Original: {elem.original}\n")
                    if elem.translation: buffer.write(f"  │ 🇸🇦 Translation: {elem.translation}\n")
                    if elem.alternative: buffer.write(f"  │ 🔄 Alternative: {elem.alternative}\n")
                    if elem.reason: buffer.write(f"  │ 💡 Reason: {elem.reason}\n")
                    buffer.write("  └──────────────────────────────────────────\n")
                buffer.write("\n")
        val = buffer.getvalue()
        buffer.close()
        return io.BytesIO(val.encode('utf-8'))

    @staticmethod
    def generate_docx(page_data_list: List[PageData]) -> io.BytesIO:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        for page_idx, page_data in enumerate(page_data_list, 1):
            file_name = page_data.file_name or "Unknown"
            p_title = doc.add_paragraph()
            run_title = p_title.add_run(f'📄 Page {page_idx}')
            run_title.bold = True
            run_title.font.size = Pt(18)
            run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            
            p_file = doc.add_paragraph()
            run_file = p_file.add_run(f'🖼️ File: {file_name}')
            run_file.italic = True
            run_file.font.size = Pt(10)
            run_file.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            
            p_div = doc.add_paragraph()
            p_div.add_run('_' * 60).font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)
            
            if not page_data or not page_data.scenes:
                doc.add_paragraph("⚠️ No data extracted.")
                doc.add_page_break()
                continue
            
            for scene in page_data.scenes:
                env = scene.environment or "Unknown Environment"
                h = doc.add_heading(f'🎬 Scene {scene.scene_number}: {env}', level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x2E, 0x74, 0x05)
                for elem in scene.elements:
                    speaker = elem.speaker or "N/A"
                    elem_type = elem.type or "text"
                    p_elem = doc.add_paragraph()
                    p_elem.paragraph_format.left_indent = Pt(20)
                    run_num = p_elem.add_run(f'Element {elem.element_number} ')
                    run_num.bold = True
                    run_num.font.size = Pt(12)
                    run_type = p_elem.add_run(f'({elem_type})')
                    run_type.italic = True
                    run_type.font.size = Pt(9)
                    run_type.font.color.rgb = RGBColor(0xA6, 0xA6, 0xA6)
                    run_spk = p_elem.add_run(f' - 🗣 {speaker}')
                    run_spk.bold = True
                    run_spk.font.size = Pt(11)
                    run_spk.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    if elem.original:
                        p_orig = doc.add_paragraph()
                        p_orig.paragraph_format.left_indent = Pt(40)
                        run_orig_lbl = p_orig.add_run('📝 Original: ')
                        run_orig_lbl.bold = True
                        run_orig_lbl.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                        p_orig.add_run(elem.original)
                    if elem.translation:
                        p_trans = doc.add_paragraph()
                        p_trans.paragraph_format.left_indent = Pt(40)
                        run_trans_lbl = p_trans.add_run('🇸🇦 Translation: ')
                        run_trans_lbl.bold = True
                        run_trans_lbl.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                        run_trans_text = p_trans.add_run(elem.translation)
                        run_trans_text.font.size = Pt(12)
                        run_trans_text.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    if elem.alternative:
                        p_alt = doc.add_paragraph()
                        p_alt.paragraph_format.left_indent = Pt(40)
                        run_alt_lbl = p_alt.add_run('🔄 Alternative: ')
                        run_alt_lbl.bold = True
                        p_alt.add_run(elem.alternative)
                    if elem.reason:
                        p_reason = doc.add_paragraph()
                        p_reason.paragraph_format.left_indent = Pt(40)
                        run_reason_lbl = p_reason.add_run('💡 Reason: ')
                        run_reason_lbl.bold = True
                        p_reason.add_run(elem.reason).italic = True
                    doc.add_paragraph()
                doc.add_paragraph()
            doc.add_page_break()

        byte_io = io.BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        return byte_io